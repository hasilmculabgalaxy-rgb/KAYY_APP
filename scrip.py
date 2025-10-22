

# KAY App - FINAL SINGLE PAGE APP (SPA) - VERSI DIPERBARUI & TAMPILAN LEBIH MENARIK
# - **FITUR LAMA LENGKAP:** Gabung, Pisah, Encrypt, Reorder, Kompres Foto.
# - **FITUR BARU LENGKAP:** Batch Rename PDF/Gambar Sesuai Excel/Sequential, Organise MCU by Excel.
# - **FITUR DIPERBARUI:** Dashboard Analisis Data MCU Massal.
# - **FITUR TERBARU:** Terjemahan PDF ke Bahasa Lain.

import os
import io
import zipfile
import shutil
import traceback
import tempfile
import time
import re
from typing import Dict, List, Optional, Tuple, Union

import streamlit as st
import pandas as pd
from PIL import Image

# PDF libs
try:
    from PyPDF2 import PdfReader, PdfWriter
except ImportError:
    PdfReader = PdfWriter = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

# pdf2image: try both convert_from_bytes and convert_from_path availability
PDF2IMAGE_AVAILABLE = False
convert_from_bytes = convert_from_path = None
try:
    from pdf2image import convert_from_bytes, convert_from_path 
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    try:
        from pdf2image import convert_from_path
        PDF2IMAGE_AVAILABLE = True
        convert_from_bytes = None
    except ImportError:
        pass

# New imports for translation
try:
    from deep_translator import GoogleTranslator
    Translator = GoogleTranslator
except ImportError:
    Translator = None

# ----------------- Constants -----------------
APP_TITLE = "Master - App – Tools MCU"
APP_DESCRIPTION = "Aplikasi serbaguna untuk pengolahan dokumen, PDF, gambar, dan data MCU — UI elegan + fungsi lengkap."
PAGE_ICON = "🧰"
SUPPORTED_IMAGE_TYPES = ["jpg", "jpeg", "png", "webp"]
SUPPORTED_PDF_TYPES = ["pdf"]
SUPPORTED_EXCEL_TYPES = ["xlsx", "csv"]
SUPPORTED_DOC_TYPES = ["docx", "doc"]

# ----------------- Helper Functions -----------------
def initialize_session_state():
    """Initialize session state variables"""
    if "menu_selection" not in st.session_state:
        st.session_state.menu_selection = "Dashboard"

def make_zip_from_map(bytes_map: Dict[str, bytes]) -> bytes:
    """Create a zip file from a dictionary of file names and bytes"""
    b = io.BytesIO()
    with zipfile.ZipFile(b, "w") as z:
        for name, data in bytes_map.items():
            z.writestr(name, data)
    b.seek(0)
    return b.getvalue()

def df_to_excel_bytes(df: pd.DataFrame) -> bytes:
    """Convert a DataFrame to Excel bytes"""
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer: 
        df.to_excel(writer, index=False)
    out.seek(0)
    return out.getvalue()

def try_encrypt(writer: PdfWriter, password: str):
    """Try to encrypt a PDF with the given password"""
    try:
        writer.encrypt(password)
    except TypeError:
        try:
            writer.encrypt(user_pwd=password, owner_pwd=None)
        except Exception:
            writer.encrypt(user_pwd=password, owner_pwd=password)

def rotate_page_safe(page, angle: int):
    """Safely rotate a PDF page"""
    try:
        page.rotate(angle)
    except Exception:
        try:
            from PyPDF2.generic import NameObject, NumberObject
            page.__setitem__(NameObject("/Rotate"), NumberObject(angle))
        except Exception:
            pass
          
def navigate_to(target_menu: str):
    """Navigate to a specific menu"""
    st.session_state.menu_selection = target_menu
    try:
        st.rerun() 
    except AttributeError:
        st.experimental_rerun()

def add_back_to_dashboard_button():
    """Add a back to dashboard button"""
    if st.button("🔙 Kembali ke Dashboard", key="back_to_dash"):
        navigate_to("Dashboard")
    st.markdown("---")

def check_library_availability() -> Dict[str, bool]:
    """Check if required libraries are available"""
    return {
        "PyPDF2": PdfReader is not None,
        "pdfplumber": pdfplumber is not None,
        "python-docx (Document)": Document is not None,
        "pdf2image (convert_from_path/bytes)": PDF2IMAGE_AVAILABLE,
        "deep_translator (GoogleTranslator)": Translator is not None, 
    }

def display_library_status():
    """Display the status of required libraries"""
    st.markdown("---")
    st.subheader("Status Library Tambahan")
    st.info("Fitur ini membantu Anda mengecek apakah library Python yang dibutuhkan sudah terinstall di lingkungan Streamlit ini.")
    
    libs = check_library_availability()
    for name, is_available in libs.items():
        status = "✅ Tersedia" if is_available else "❌ Tidak Tersedia"
        st.markdown(f"- **{name}**: {status}")

def clean_column_names(df: pd.DataFrame) -> pd.DataFrame:
    """Clean column names by removing special characters and converting to lowercase"""
    df.columns = df.columns.str.replace('[^A-Za-z0-9_]+', '', regex=True).str.lower()
    return df

def preprocess_text_for_layout(text_blocks: List[str]) -> List[str]:
    """Preprocess text blocks to improve layout in translated documents"""
    processed_blocks = []
    current_block = ""
    MAX_LINE_LENGTH = 100
    
    for block in text_blocks:
        if block.strip() == "":
            if current_block:
                processed_blocks.append(current_block)
            processed_blocks.append("")
            current_block = ""
            continue
        
        if len(block.strip()) < MAX_LINE_LENGTH and not block.strip().endswith('.'):
            if current_block:
                current_block += " | " + block.strip()
            else:
                current_block = block.strip()
        else:
            if current_block:
                processed_blocks.append(current_block)
            current_block = block.strip()
            processed_blocks.append(current_block)
            current_block = ""
            
    if current_block:
        processed_blocks.append(current_block)
    
    return [b.replace("  ", " ").strip() for b in processed_blocks if b.strip() or b == ""]

def clean_folder_name(name: str) -> str:
    """Clean folder name by replacing invalid characters"""
    return str(name).strip().replace('/', '_').replace('\\', '_')

# ----------------- Streamlit config & CSS -----------------
st.set_page_config(
    page_title=APP_TITLE, 
    page_icon=PAGE_ICON, 
    layout="wide", 
    initial_sidebar_state="collapsed"
)

# CSS / Theme
st.markdown("""
<style>
/* 1. HILANGKAN SEMUA UI SIDEBAR */
[data-testid="stSidebarToggleButton"], 
section[data-testid="stSidebar"],      
[data-testid="stDecoration"]        
{
    visibility: hidden !important;
    display: none !important;
    width: 0 !important;
    padding: 0 !important;
}

/* 2. Hilangkan logo GitHub 'Fork' */
.stApp a[href*="github.com/"],
.stApp header > div:last-child {
    display: none !important;
}

/* 3. Background gradient lembut */
.stApp {
    background: linear-gradient(180deg, #e9f2ff 0%, #f4f9ff 100%); 
    color: #002b5b;
    font-family: 'Inter', sans-serif;
}

/* 4. Tombol modern glossy */
div.stButton > button {
    background: linear-gradient(90deg, #5dade2, #3498db);
    color: white; 
    border: none;
    border-radius: 12px;
    padding: 0.7rem 1.2rem;
    font-weight: 600;
    transition: 0.2s;
    box-shadow: 0 4px 10px rgba(52, 152, 219, 0.4); 
    cursor: pointer;
    width: auto;
}

div.stButton > button[key*="dash_"], 
div.stButton > button[key*="back_"] {
    width: 100%;
    margin-top: 10px; 
}

div.stButton > button:hover {
    background: linear-gradient(90deg, #3498db, #2e86c1); 
    transform: scale(1.02);
    box-shadow: 0 8px 18px rgba(52, 152, 219, 0.6);
}

/* 5. Card fitur */
.feature-card {
    background: white;
    border-radius: 16px;
    box-shadow: 0 4px 15px rgba(0,0,0,0.08);
    padding: 24px;
    transition: all 0.3s ease-in-out;
    border: 1px solid #d0e3ff; 
    display: flex;
    flex-direction: column;
    justify-content: space-between;
    height: 100%; 
}
.feature-card:hover {
    box-shadow: 0 10px 30px rgba(0,0,0,0.15);
    transform: translateY(-5px);
    border-color: #3498db;
}

/* 6. Small UI tweaks */
h1 { color: #1b4f72; font-weight: 800; }
h2, h3, h4 { color: #2e86c1; }
.stInfo, .stWarning {
    border-radius: 12px;
    padding: 1rem;
}
.stInfo { background-color: #e3f2fd; border-left: 5px solid #2196f3; }
.stWarning { background-color: #fff3e0; border-left: 5px solid #ff9800; }

/* 7. Kontainer untuk merapikan section */
.dashboard-container {
    padding: 20px 0;
    margin-bottom: 20px;
}
</style>
""", unsafe_allow_html=True)

# ----------------- Global Header & Navigation Setup -----------------
initialize_session_state()
menu = st.session_state.menu_selection

# ----------------- Halaman Header -----------------
header_col1, header_col2 = st.columns([1, 4])
with header_col1:
    st.markdown("<h1 style='font-size: 3rem;'>🧰</h1>", unsafe_allow_html=True)

with header_col2:
    st.title(APP_TITLE)
    st.markdown(APP_DESCRIPTION)

st.markdown("---")

# -----------------------------------------------------------------------------
# ----------------- FUNGSI UTAMA -----------------
# -----------------------------------------------------------------------------

# -------------- Dashboard (Menu Utama) --------------
def render_dashboard():
    st.markdown('<div class="dashboard-container">', unsafe_allow_html=True)
    st.markdown("## 🚀 Pilih Fitur Utama")
    
    # ------------------ FITUR UTAMA ------------------
    cols1 = st.columns(3)

    # Kompres Foto
    with cols1[0]:
        with st.container():
            st.markdown('<div class="feature-card"><b>🖼️ Kompres Foto / Gambar Tools</b><br>Perkecil ukuran, ubah format, Batch Rename Sesuai Excel.</div>', unsafe_allow_html=True)
            if st.button("Buka Kompres Foto", key="dash_foto"):
                navigate_to("Kompres Foto")

    # PDF Tools
    with cols1[1]:
        with st.container():
            st.markdown('<div class="feature-card"><b>📄 PDF Tools</b><br>Gabung, pisah, encrypt, Reorder & Batch Rename Sesuai Excel/Sequential, **Terjemahan**.</div>', unsafe_allow_html=True)
            if st.button("Buka PDF Tools", key="dash_pdf"):
                navigate_to("PDF Tools")

    # MCU Tools
    with cols1[2]:
        with st.container():
            st.markdown('<div class="feature-card"><b>🩺 MCU Tools</b><br>Proses Excel + PDF untuk hasil MCU / Analisis Data. **Termasuk Organise by Excel**</div>', unsafe_allow_html=True)
            if st.button("Buka MCU Tools", key="dash_mcu"):
                navigate_to("MCU Tools")
            
    # ------------------ FITUR LAINNYA ------------------
    st.markdown("## 🛠️ Fitur Lainnya")
    cols2 = st.columns(3)
    
    # File Tools
    with cols2[0]:
        with st.container():
            st.markdown('<div class="feature-card"><b>📁 File Tools</b><br>Zip/unzip, konversi dasar, Batch Rename Gambar & PDF.</div>', unsafe_allow_html=True)
            if st.button("Buka File Tools", key="dash_file"):
                navigate_to("File Tools")

    # Tentang
    with cols2[1]:
        with st.container():
            st.markdown('<div class="feature-card"><b>💡 Tentang Aplikasi</b><br>Informasi dan kebutuhan library.</div>', unsafe_allow_html=True)
            if st.button("Lihat Tentang", key="dash_about"):
                navigate_to("Tentang")

    # Kolom kosong
    with cols2[2]:
        st.markdown('<div class="feature-card" style="visibility:hidden; height: 100%;">.</div>', unsafe_allow_html=True)
        
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.info("Semua proses berlangsung lokal di perangkat server tempat Streamlit dijalankan.")

# -------------- Kompres Foto / Image Tools --------------
def render_compress_foto():
    add_back_to_dashboard_button() 
    st.subheader("🖼️ Kompres & Kelola Foto/Gambar")
    
    img_tool = st.selectbox("Pilih Fitur Gambar", [
        "🖼️ Kompres Foto (Batch)", 
        "📝 Batch Rename/Format Gambar (Sequential)",
        "📊 Batch Rename Gambar Sesuai Excel (Fitur Baru)"
        ])

    if img_tool == "🖼️ Kompres Foto (Batch)":
        render_compress_foto_batch()
    elif img_tool == "📝 Batch Rename/Format Gambar (Sequential)":
        render_batch_rename_sequential_images()
    elif img_tool == "📊 Batch Rename Gambar Sesuai Excel (Fitur Baru)":
        render_batch_rename_excel_images()

def render_compress_foto_batch():
    st.markdown("---")
    st.markdown("### 🖼️ Kompres Foto (Batch)")
    uploaded = st.file_uploader("Unggah gambar (jpg/png) — bisa banyak", type=["jpg","jpeg","png"], accept_multiple_files=True)
    quality = st.slider("Kualitas JPEG", 10, 95, 75)
    max_side = st.number_input("Max side (px)", min_value=100, max_value=4000, value=1200)
    
    if uploaded and st.button("Kompres Semua"):
        out_map = {}
        total = len(uploaded)
        prog = st.progress(0)
        
        with st.spinner("Mengompres..."):
            for i, f in enumerate(uploaded):
                try:
                    raw = f.read()
                    im = Image.open(io.BytesIO(raw))
                    im.thumbnail((max_side, max_side))
                    buf = io.BytesIO()
                    im.convert("RGB").save(buf, format="JPEG", quality=quality, optimize=True)
                    out_map[f"compressed_{f.name}"] = buf.getvalue()
                except Exception as e:
                    st.warning(f"Gagal: {f.name} — {e}")
                prog.progress(int((i+1)/total*100))
        
        if out_map:
            zipb = make_zip_from_map(out_map)
            st.success(f"✅ {len(out_map)} file berhasil dikompres")
            st.download_button("Unduh Hasil (ZIP)", zipb, file_name="foto_kompres.zip", mime="application/zip")
        else:
            st.warning("Tidak ada file berhasil dikompres.")

def render_batch_rename_sequential_images():
    st.markdown("---")
    st.markdown("### 📝 Ganti Nama & Ubah Format Gambar Massal (Sequential)")
    uploaded_files = st.file_uploader(
        "Unggah file Gambar (JPG, PNG, dll.):", 
        type=["jpg", "jpeg", "png", "webp"], 
        accept_multiple_files=True,
        key="batch_rename_uploader"
    )
    
    if uploaded_files:
        col1, col2 = st.columns(2)
        new_prefix = col1.text_input("Prefix Nama File Baru:", value="KAY_File", help="Contoh: KAY_File_001.jpg", key="prefix_img_seq")
        new_format = col2.selectbox("Format Output Baru:", ["Sama seperti Asli", "JPG", "PNG", "WEBP"], index=0, key="format_img_seq")

        if st.button("Proses Batch File", key="process_batch_rename_seq"):
            if not new_prefix: 
                st.error("Prefix nama file tidak boleh kosong.")
                st.stop()
            else:
                output_zip = io.BytesIO()
                try:
                    with zipfile.ZipFile(output_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for i, file in enumerate(uploaded_files, 1):
                            _, original_ext = os.path.splitext(file.name)
                            img = Image.open(file)
                            img_io = io.BytesIO()
                            
                            if new_format == "Sama seperti Asli":
                                output_format_pil = img.format if img.format else 'JPEG'
                                output_ext = original_ext
                            else:
                                output_ext = "." + new_format.lower()
                                output_format_pil = new_format.upper()
                                
                            new_filename = f"{new_prefix}_{i:03d}{output_ext}"
                            
                            if output_format_pil in ('JPEG', 'JPG'):
                                img.convert("RGB").save(img_io, format='JPEG', quality=95) 
                            elif output_format_pil == 'PNG':
                                img.save(img_io, format='PNG')
                            elif output_format_pil == 'WEBP':
                                img.save(img_io, format='WEBP')
                            else:
                                img.save(img_io, format=output_format_pil) 
                                
                            img_io.seek(0)
                            zf.writestr(new_filename, img_io.read())
                    
                    st.success(f"✅ Berhasil memproses {len(uploaded_files)} file.") 
                    st.download_button("Unduh File ZIP Hasil Batch", data=output_zip.getvalue(), file_name="hasil_batch_gambar.zip", mime="application/zip")
                except Exception as e: 
                    st.error(f"Gagal memproses file: {e}")
                    traceback.print_exc()

def render_batch_rename_excel_images():
    st.markdown("---")
    st.markdown("### 📊 Ganti Nama Gambar (PNG/JPEG) Berdasarkan Excel")
    st.info("Template Excel/CSV wajib memiliki kolom **`nama_lama`** (termasuk ekstensi, misal: `foto_123.jpg`) dan **`nama_baru`** (termasuk ekstensi, misal: `ID_001.png`).")
    
    excel_up = st.file_uploader("Unggah Excel/CSV untuk daftar nama:", type=["xlsx", "csv"], key="rename_img_excel_up")
    files = st.file_uploader("Unggah Gambar (JPG/PNG/JPEG, multiple):", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="rename_img_files_up")
  
    if excel_up and files and st.button("Proses Ganti Nama Gambar (ZIP)", key="process_img_rename_excel"):
        try:
            with st.spinner("Memproses penggantian nama..."):
                # 1. Baca Excel
                if excel_up.name.lower().endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(excel_up.read()))
                else:
                    df = pd.read_excel(io.BytesIO(excel_up.read()))
                
                # 2. Validasi Kolom
                required_cols = ['nama_lama', 'nama_baru']
                if not all(col in df.columns for col in required_cols):
                    st.error(f"Excel/CSV wajib memiliki kolom: {', '.join(required_cols)}")
                    st.stop()
                
                # 3. Map File dan Proses Rename
                file_map = {f.name: f.read() for f in files}
                out_map = {}
                not_found = []
                df['nama_lama_str'] = df['nama_lama'].astype(str).str.strip()

                for _, row in df.iterrows():
                    old_name = str(row['nama_lama']).strip()
                    new_name = str(row['nama_baru']).strip()
                    
                    if old_name in file_map:
                        if not os.path.splitext(new_name)[1]:
                            _, old_ext = os.path.splitext(old_name)
                            new_name = new_name + old_ext
                        out_map[new_name] = file_map[old_name]
                    else:
                        not_found.append(old_name)

                # 4. Buat ZIP
                if out_map:
                    zipb = make_zip_from_map(out_map)
                    st.success(f"✅ {len(out_map)} file berhasil diganti namanya dan dikemas.") 
                    st.download_button("Unduh Hasil (ZIP)", zipb, file_name="gambar_renamed_by_excel.zip", mime="application/zip")
                else:
                    st.warning("Tidak ada file yang cocok ditemukan atau diproses.")
                
                if not_found:
                    st.info(f"{len(not_found)} file 'nama_lama' di Excel tidak ditemukan di file yang diunggah. Contoh: {not_found[:5]}")
        except Exception as e:
            st.error(f"Terjadi kesalahan pemrosesan: {e}")
            traceback.print_exc()

# -------------- PDF Tools --------------
def render_pdf_tools():
    add_back_to_dashboard_button() 
    st.subheader("📄 PDF Tools")

    pdf_options = [
        "--- Pilih Tools ---",
        "🔗 Gabung PDF", 
        "✂️ Pisah PDF", 
        "🔄 Reorder/Hapus Halaman PDF", 
        "📝 Batch Rename PDF (Sequential)", 
        "📊 Batch Rename PDF Sesuai Excel (Fitur Baru)", 
        "🖼️ Image -> PDF", 
        "🖼️ PDF -> Image", 
        "📄 Ekstraksi Teks/Tabel", 
        "🌐 Terjemahan PDF ke Bahasa Lain (Fitur Baru)",
        "🔄 Konversi PDF", 
        "🔒 Proteksi PDF", 
        "🛠️ Utility PDF", 
    ]
    
    tool_select = st.selectbox("Pilih fitur PDF", pdf_options)

    # Mapping
    if tool_select == "--- Pilih Tools ---": 
        tool = None
    elif tool_select == "📄 Ekstraksi Teks/Tabel": 
        tool = st.selectbox("Pilih mode ekstraksi", ["Extract Text", "Extract Tables -> Excel"]) 
    elif tool_select == "🔄 Konversi PDF": 
        tool = st.selectbox("Pilih mode konversi", ["PDF -> Word", "PDF -> Excel (text)"]) 
    elif tool_select == "🔒 Proteksi PDF": 
        tool = st.selectbox("Pilih mode proteksi", ["Encrypt PDF", "Decrypt PDF", "Batch Lock (Excel)"]) 
    elif tool_select == "🛠️ Utility PDF": 
        tool = st.selectbox("Pilih mode utilitas", ["Hapus Halaman", "Rotate PDF", "Kompres PDF", "Watermark PDF", "Preview PDF"]) 
    elif tool_select == "🔗 Gabung PDF": 
        tool = "Gabung PDF" 
    elif tool_select == "✂️ Pisah PDF": 
        tool = "Pisah PDF" 
    elif tool_select == "🔄 Reorder/Hapus Halaman PDF": 
        tool = "Reorder PDF" 
    elif tool_select == "📝 Batch Rename PDF (Sequential)": 
        tool = "Batch Rename PDF Seq" 
    elif tool_select == "📊 Batch Rename PDF Sesuai Excel (Fitur Baru)": 
        tool = "Batch Rename PDF Excel" 
    elif tool_select == "🖼️ PDF -> Image": 
        tool = "PDF -> Image" 
    elif tool_select == "🖼️ Image -> PDF": 
        tool = "Image -> PDF" 
    elif tool_select == "🌐 Terjemahan PDF ke Bahasa Lain (Fitur Baru)": 
        tool = "Translate PDF"
    else: 
        tool = None
    
    # Render the selected tool
    if tool == "Translate PDF":
        render_translate_pdf()
    elif tool == "Batch Rename PDF Excel":
        render_batch_rename_excel_pdf()
    elif tool == "Batch Rename PDF Seq":
        render_batch_rename_sequential_pdf()
    elif tool == "Reorder PDF":
        render_reorder_pdf()
    elif tool == "Gabung PDF":
        render_merge_pdf()
    elif tool == "Pisah PDF":
        render_split_pdf()
    elif tool == "Hapus Halaman":
        render_delete_page()
    elif tool == "Rotate PDF":
        render_rotate_pdf()
    elif tool == "Kompres PDF":
        render_compress_pdf()
    elif tool == "Watermark PDF":
        render_watermark_pdf()
    elif tool == "PDF -> Image":
        render_pdf_to_image()
    elif tool == "Image -> PDF":
        render_image_to_pdf()
    elif tool == "Extract Text":
        render_extract_text()
    elif tool == "Extract Tables -> Excel":
        render_extract_tables()
    elif tool == "PDF -> Word":
        render_pdf_to_word()
    elif tool == "PDF -> Excel (text)":
        render_pdf_to_excel()
    elif tool == "Encrypt PDF":
        render_encrypt_pdf()
    elif tool == "Decrypt PDF":
        render_decrypt_pdf()
    elif tool == "Batch Lock (Excel)":
        render_batch_lock()
    elif tool == "Preview PDF":
        render_preview_pdf()

def render_translate_pdf():
    st.markdown("---")
    st.markdown("### 🌐 Terjemahan Teks PDF (Optimasi Agar Lebih Rapi)")
    st.info("Fitur ini mencoba membuat hasil Word lebih rapi dengan menggabungkan baris-baris pendek yang berdekatan (*pre-processing*). **Replikasi tata letak kolom/tabel PDF tetap terbatas.**")
    
    if Translator is None or Document is None:
        if Translator is None: st.error("Library `deep-translator` tidak ditemukan.")
        if Document is None: st.error("Library `python-docx` tidak ditemukan.")
        st.stop()
    
    f = st.file_uploader("Unggah PDF untuk Diterjemahkan:", type="pdf", key="translate_pdf_uploader")
    
    col1, col2 = st.columns(2)
    src_lang = col1.text_input("Bahasa Sumber (ISO Code, ex: id)", value="auto", help="Ketik 'auto' jika tidak yakin.")
    target_lang = col2.text_input("Bahasa Tujuan (ISO Code, ex: en, ja, fr)", value="en")

    if f and st.button("Proses Terjemahan dan Buat Word (.docx)", key="translate_pdf_button"):
        try:
            # 1. Ekstraksi Teks
            with st.spinner("1. Mengekstrak dan merapikan teks dari PDF..."):
                raw = f.read()
                all_text_lines = []
                
                if pdfplumber:
                    with pdfplumber.open(io.BytesIO(raw)) as doc:
                        for p in doc.pages:
                            page_text = p.extract_text() or ""
                            all_text_lines.extend(page_text.split('\n'))
                            all_text_lines.append("---HALAMAN BARU---")
                else:
                    reader = PdfReader(io.BytesIO(raw))
                    for p in reader.pages:
                        page_text = p.extract_text() or ""
                        all_text_lines.extend(page_text.split('\n'))
                        all_text_lines.append("---HALAMAN BARU---")

                preprocessed_paragraphs = preprocess_text_for_layout(all_text_lines)
                full_text_clean = "\n\n".join(p for p in preprocessed_paragraphs if p != "---HALAMAN BARU---" and p != "")

            if not full_text_clean.strip():
                st.warning("Teks kosong atau tidak dapat diekstrak dari PDF.")
                st.stop()

            # 2. Chunking Berbasis Paragraf dan Terjemahan
            with st.spinner(f"2. Menerjemahkan teks ke {target_lang} (mempertahankan paragraf)..."):
                translator = Translator(source=src_lang, target=target_lang)
                CHUNK_SIZE = 4500
                
                current_chunk = ""
                text_chunks_for_translation = []
                
                for p in preprocessed_paragraphs:
                    if p == "---HALAMAN BARU---":
                        if current_chunk:
                            text_chunks_for_translation.append(current_chunk)
                        text_chunks_for_translation.append("---HALAMAN BARU---")
                        current_chunk = ""
                        continue
                    
                    if not p.strip(): 
                         if current_chunk:
                            text_chunks_for_translation.append(current_chunk)
                         text_chunks_for_translation.append("")
                         current_chunk = ""
                         continue

                    if len(current_chunk) + len(p) + 4 > CHUNK_SIZE:
                        if current_chunk:
                            text_chunks_for_translation.append(current_chunk)
                        current_chunk = p + "\n\n"
                    else:
                        current_chunk += p + "\n\n"
                    
                    if current_chunk:
                        text_chunks_for_translation.append(current_chunk.strip())
                        
                translated_parts = []
                prog = st.progress(0)
                
                for i, chunk in enumerate(text_chunks_for_translation):
                    if chunk in ("---HALAMAN BARU---", ""):
                        translated_parts.append(chunk) 
                    else:
                        if i > 0: time.sleep(0.1) 
                        translated = translator.translate(chunk)
                        translated_parts.append(translated.replace(" | ", " | ").strip()) 
                    
                    prog.progress(int((i + 1) / len(text_chunks_for_translation) * 100))

                translated_text_combined = "\n\n".join(translated_parts)
                prog.empty()

            # 3. Rekonstruksi ke Word
            with st.spinner("3. Membuat file Word (.docx) baru..."):
                doc = Document()
                
                for item in translated_text_combined.split('\n\n'):
                    item_stripped = item.strip()

                    if item_stripped == "---HALAMAN BARU---":
                        doc.add_page_break()
                    elif item_stripped:
                        doc.add_paragraph(item_stripped)
                
                out = io.BytesIO()
                doc.save(out)
                out.seek(0)
                
            st.success("✅ Terjemahan berhasil! Unduh file Word hasil terjemahan.")
            st.download_button(
                f"📥 Unduh Hasil Terjemahan ({target_lang}).docx", 
                data=out.getvalue(), 
                file_name=f"translated_to_{target_lang}_rapi.docx", 
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            st.markdown("---")
            st.markdown("#### Preview Teks Asli (Sudah Dirapikan) dan Terjemahan")
            col_preview1, col_preview2 = st.columns(2)
            with col_preview1:
                st.text_area("Teks Asli (Dirapikan)", "\n\n".join(preprocessed_paragraphs)[:5000], height=300)
            with col_preview2:
                st.text_area("Teks Terjemahan", translated_text_combined[:5000], height=300)

        except Exception as e:
            st.error(f"Terjadi kesalahan saat terjemahan. Cek kode bahasa (ISO 639-1) dan pastikan teks dapat diekstrak. Error: {e}")
            traceback.print_exc()

def render_batch_rename_excel_pdf():
    st.markdown("---")
    st.markdown("### 📊 Ganti Nama File PDF Berdasarkan Excel") 
    st.markdown("Unggah banyak file PDF dan ganti namanya sesuai daftar di Excel/CSV.")
    st.info("Template Excel/CSV wajib memiliki kolom **`nama_lama`** (misal: `ID_123.pdf`) dan **`nama_baru`** (misal: `Hasil_123.pdf`).")

    excel_up = st.file_uploader("Unggah Excel/CSV untuk daftar nama:", type=["xlsx", "csv"], key="rename_pdf_excel_up")
    files = st.file_uploader("Unggah File PDF (multiple):", type=["pdf"], accept_multiple_files=True, key="rename_pdf_files_up")
    
    if excel_up and files and st.button("Proses Ganti Nama PDF (ZIP)", key="process_pdf_rename_excel"):
        try:
            with st.spinner("Memproses penggantian nama..."):
                # 1. Baca Excel
                if excel_up.name.lower().endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(excel_up.read()))
                else:
                    df = pd.read_excel(io.BytesIO(excel_up.read()))
                
                # 2. Validasi Kolom
                required_cols = ['nama_lama', 'nama_baru']
                if not all(col in df.columns for col in required_cols):
                    st.error(f"Excel/CSV wajib memiliki kolom: {', '.join(required_cols)}")
                    st.stop()
                
                # 3. Map File dan Proses Rename
                file_map = {f.name: f.read() for f in files}
                out_map = {}
                not_found = []
                
                for _, row in df.iterrows():
                    old_name = str(row['nama_lama']).strip()
                    new_name = str(row['nama_baru']).strip()
                    
                    if old_name in file_map:
                        if not new_name.lower().endswith('.pdf'):
                            new_name += '.pdf'
                        out_map[new_name] = file_map[old_name]
                    else:
                        not_found.append(old_name)

                # 4. Buat ZIP
                if out_map:
                    zipb = make_zip_from_map(out_map)
                    st.success(f"✅ {len(out_map)} file berhasil diganti namanya dan dikemas.") 
                    st.download_button("Unduh Hasil (ZIP)", zipb, file_name="pdf_renamed_by_excel.zip", mime="application/zip")
                else:
                    st.warning("Tidak ada file yang cocok ditemukan atau diproses.")
                
                if not_found:
                    st.info(f"{len(not_found)} file 'nama_lama' di Excel tidak ditemukan di file yang diunggah. Contoh: {not_found[:5]}")
        except Exception as e:
            st.error(f"Terjadi kesalahan pemrosesan: {e}")
            traceback.print_exc()

def render_batch_rename_sequential_pdf():
    st.markdown("---")
    st.markdown("### 📝 Ganti Nama File PDF Massal (Sequential)") 
    uploaded_files = st.file_uploader("Unggah file PDF (multiple):", type=["pdf"], accept_multiple_files=True, key="batch_rename_pdf_uploader_seq")

    if uploaded_files:
        col1, col2 = st.columns(2)
        new_prefix = col1.text_input("Prefix Nama File Baru:", value="Hasil_PDF", help="Contoh: Hasil_PDF_001.pdf", key="prefix_pdf_seq")
        start_num = col2.number_input("Mulai dari Angka (Counter):", min_value=1, value=1, step=1, key="start_num_pdf_seq")

        if st.button("Proses Ganti Nama (ZIP)", key="process_batch_rename_pdf_seq"):
            if not new_prefix: 
                st.error("Prefix nama file tidak boleh kosong.")
                st.stop()
            else:
                output_zip = io.BytesIO()
                try:
                    with zipfile.ZipFile(output_zip, 'w', zipfile.ZIP_DEFLATED) as zf:
                        for i, file in enumerate(uploaded_files, start_num):
                            new_filename = f"{new_prefix}_{i:03d}.pdf"
                            zf.writestr(new_filename, file.read())
                    st.success(f"✅ Berhasil mengganti nama {len(uploaded_files)} file.") 
                    st.download_button("Unduh File ZIP Hasil Rename", data=output_zip.getvalue(), file_name="pdf_renamed.zip", mime="application/zip")
                except Exception as e: 
                    st.error(f"Gagal memproses file: {e}")
                    traceback.print_exc()

def render_reorder_pdf():
    st.markdown("---")
    st.markdown("### 🔄 Reorder atau Hapus Halaman PDF") 
    st.markdown("Unggah file PDF Anda dan tentukan urutan halaman baru (contoh: `2, 1, 3` untuk membalik, atau `1, 3` untuk menghapus halaman 2).")

    f = st.file_uploader("Unggah 1 file PDF:", type="pdf", key="reorder_pdf_uploader")
    
    if f:
        try:
            raw = f.read()
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            reader = PdfReader(io.BytesIO(raw))
            num_pages = len(reader.pages)
            st.info(f"PDF berhasil dimuat. Jumlah total halaman: **{num_pages}**.")
            
            default_order = ", ".join(map(str, range(1, num_pages + 1)))

            new_order_str = st.text_input(
                f"Masukkan urutan halaman baru (1-{num_pages}) dipisahkan koma:",
                value=default_order,
                help="Contoh: '3, 1, 2' untuk mengubah urutan. '1, 3, 5' untuk menghapus halaman genap."
            )

            if st.button("Proses Reorder/Hapus Halaman", key="process_reorder"):
                new_order_indices = []
                
                try:
                    input_list = [int(x.strip()) for x in new_order_str.split(',') if x.strip().isdigit()]
                    
                    if any(n < 1 or n > num_pages for n in input_list):
                        st.error(f"Nomor halaman harus antara 1 sampai {num_pages}.")
                        st.stop() 

                    new_order_indices = [n - 1 for n in input_list]
                  
                    writer = PdfWriter()
                    for index in new_order_indices:
                        writer.add_page(reader.pages[index])
                    
                    pdf_buffer = io.BytesIO()
                    writer.write(pdf_buffer)
                    pdf_buffer.seek(0)

                    st.download_button(
                        "📥 Unduh Hasil PDF (Reordered)", 
                        data=pdf_buffer,
                        file_name="pdf_reordered.pdf",
                        mime="application/pdf"
                    )
                    st.success(f"✅ Pemrosesan selesai. Total halaman baru: {len(new_order_indices)}.")

                except Exception as e:
                    st.error(f"Format urutan halaman tidak valid atau terjadi kesalahan pemrosesan: {e}")

        except Exception as e:
            st.error(f"Terjadi kesalahan saat memproses PDF: {e}")
            st.info("Pastikan file yang diunggah adalah PDF yang valid.")

def render_merge_pdf():
    st.markdown("---")
    st.markdown("### 🔗 Gabung PDF")
    files = st.file_uploader("Upload PDFs (multiple):", type="pdf", accept_multiple_files=True)
    if files and st.button("Gabung"):
        try:
            if PdfWriter is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Menggabungkan..."):
                writer = PdfWriter()
                for f in files:
                    reader = PdfReader(io.BytesIO(f.read()))
                    for p in reader.pages:
                        writer.add_page(p)
                out = io.BytesIO(); writer.write(out); out.seek(0)
            st.download_button("📥 Download merged.pdf", out.getvalue(), file_name="merged.pdf", mime="application/pdf")
            st.success("✅ Selesai")
        except Exception:
            st.error(traceback.format_exc())

def render_split_pdf():
    st.markdown("---")
    st.markdown("### ✂️ Pisah PDF")
    f = st.file_uploader("Upload single PDF:", type="pdf")
    if f and st.button("Split to pages (ZIP)"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Memisahkan..."):
                reader = PdfReader(io.BytesIO(f.read()))
                out_map = {}
                for i, p in enumerate(reader.pages):
                    w = PdfWriter(); w.add_page(p)
                    buf = io.BytesIO(); w.write(buf); buf.seek(0)
                    out_map[f"page_{i+1}.pdf"] = buf.getvalue()
                zipb = make_zip_from_map(out_map)
            st.download_button("📥 Download pages.zip", zipb, file_name="pages.zip", mime="application/zip")
        except Exception:
            st.error(traceback.format_exc())

def render_delete_page():
    st.markdown("---")
    st.markdown("### 🗑️ Hapus Halaman dari PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    page_no = st.number_input("Halaman yang dihapus (1-based)", min_value=1, value=1)
    if f and st.button("Hapus Halaman"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Menghapus..."):
                reader = PdfReader(io.BytesIO(f.read()))
                writer = PdfWriter()
                for i, p in enumerate(reader.pages):
                    if i+1 != page_no:
                        writer.add_page(p)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download result", buf.getvalue(), file_name="removed_page.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_rotate_pdf():
    st.markdown("---")
    st.markdown("### 🔄 Putar Halaman PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    angle = st.selectbox("Rotate degrees", [90, 180, 270])
    if f and st.button("Rotate"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Memutar..."):
                reader = PdfReader(io.BytesIO(f.read()))
                writer = PdfWriter()
                for p in reader.pages:
                    rotate_page_safe(p, angle)
                    writer.add_page(p)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download rotated.pdf", buf.getvalue(), file_name="rotated.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_compress_pdf():
    st.markdown("---")
    st.markdown("### 🗜️ Kompres Ukuran PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    if f and st.button("Compress (rewrite)"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Mengompres (rewrite)..."):
                reader = PdfReader(io.BytesIO(f.read()))
                writer = PdfWriter()
                for p in reader.pages:
                    writer.add_page(p)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download compressed.pdf", buf.getvalue(), file_name="compressed.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_watermark_pdf():
    st.markdown("---")
    st.markdown("### 💧 Tambah Watermark ke PDF")
    base = st.file_uploader("Base PDF", type="pdf")
    watermark = st.file_uploader("Watermark PDF (single page)", type="pdf")
    if base and watermark and st.button("Apply watermark"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Menerapkan watermark..."):
                rb = PdfReader(io.BytesIO(base.read()))
                rm = PdfReader(io.BytesIO(watermark.read()))
                wm = rm.pages[0]
                writer = PdfWriter()
                for p in rb.pages:
                    try:
                        p.merge_page(wm)
                    except Exception:
                        try:
                            p.mergeTranslatedPage(wm, 0, 0)
                        except Exception:
                            pass
                    writer.add_page(p)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download watermarked.pdf", buf.getvalue(), file_name="watermarked.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_pdf_to_image():
    st.markdown("---")
    st.markdown("### 🖼️ PDF ke Gambar (PNG/JPEG)")
    st.info("Requires pdf2image + poppler (server).")
    f = st.file_uploader("Upload PDF", type="pdf")
    dpi = st.slider("DPI", 100, 300, 150)
    fmt = st.radio("Format", ["PNG", "JPEG"])
    if f and st.button("Convert to images"):
        try:
            if not PDF2IMAGE_AVAILABLE:
                st.error("pdf2image not installed or poppler missing.")
                st.stop()
            else:
                with st.spinner("Converting..."):
                    pdf_bytes = f.read()
                    images = None
                    if convert_from_bytes is not None:
                        images = convert_from_bytes(pdf_bytes, dpi=dpi)
                    else:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                            tmp.write(pdf_bytes)
                        tmp_path = tmp.name
                        images = convert_from_path(tmp_path, dpi=dpi)
                        try:
                            os.unlink(tmp_path)
                        except Exception:
                            pass
                    out_map = {}
                    for i, img in enumerate(images):
                        b = io.BytesIO(); img.save(b, format=fmt); out_map[f"page_{i+1}.{fmt.lower()}"] = b.getvalue()
                    zipb = make_zip_from_map(out_map)
                    st.download_button("📥 Download images.zip", zipb, file_name="pdf_images.zip", mime="application/zip")
        except Exception:
            st.error(traceback.format_exc())

def render_image_to_pdf():
    st.markdown("---")
    st.markdown("### 🖼️ Gambar ke PDF")
    imgs = st.file_uploader("Upload images", type=["jpg","png","jpeg"], accept_multiple_files=True)
    if imgs and st.button("Images -> PDF"):
        try:
            with st.spinner("Membuat PDF dari gambar..."):
                pil = [Image.open(io.BytesIO(i.read())).convert("RGB") for i in imgs]
                buf = io.BytesIO()
                if len(pil) == 1:
                    pil[0].save(buf, format="PDF")
                else:
                    pil[0].save(buf, save_all=True, append_images=pil[1:], format="PDF")
                buf.seek(0)
            st.download_button("📥 Download images_as_pdf.pdf", buf.getvalue(), file_name="images_as_pdf.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_extract_text():
    st.markdown("---")
    st.markdown("### 📄 Ekstraksi Teks dari PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    if f and st.button("Extract text"):
        try:
            if PdfReader is None and pdfplumber is None:
                st.error("PyPDF2 atau pdfplumber tidak terinstall.")
                st.stop()
            with st.spinner("Mengekstrak teks..."):
                text_blocks = []
                raw = f.read()
                if pdfplumber:
                    with pdfplumber.open(io.BytesIO(raw)) as doc:
                        for i, p in enumerate(doc.pages):
                            text_blocks.append(f"--- Page {i+1} ---\n" + (p.extract_text() or ""))
                else:
                    reader = PdfReader(io.BytesIO(raw))
                    for i, p in enumerate(reader.pages):
                        text_blocks.append(f"--- Page {i+1} ---\n" + (p.extract_text() or ""))
                full = "\n".join(text_blocks)
                st.text_area("Extracted text (preview)", full[:10000], height=300)
                st.download_button("📥 Download .txt", full, file_name="extracted_text.txt", mime="text/plain")
        except Exception:
            st.error(traceback.format_exc())

def render_extract_tables():
    st.markdown("---")
    st.markdown("### 📊 Ekstraksi Tabel ke Excel")
    if pdfplumber is None:
        st.error("pdfplumber is required for table extraction (pip install pdfplumber)")
    else:
        f = st.file_uploader("Upload PDF", type="pdf")
        if f and st.button("Extract tables"):
            try:
                with st.spinner("Mengekstrak tabel..."):
                    tables = []
                    with pdfplumber.open(io.BytesIO(f.read())) as doc:
                        for page in doc.pages:
                            for tbl in page.extract_tables():
                                if tbl and len(tbl) > 1:
                                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                                    tables.append(df)
                    if tables:
                        df_all = pd.concat(tables, ignore_index=True)
                        st.dataframe(df_all.head())
                        excel_bytes = df_to_excel_bytes(df_all)
                        st.download_button("📥 Download Excel", data=excel_bytes, file_name="extracted_tables.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    else:
                        st.info("No tables found.")
            except Exception:
                st.error(traceback.format_exc())

def render_pdf_to_word():
    st.markdown("---")
    st.markdown("### 📄 Konversi PDF ke Word (Text-based)")
    if Document is None:
        st.error("python-docx is required for PDF->Word (pip install python-docx)")
    else:
        f = st.file_uploader("Upload PDF", type="pdf")
        if f and st.button("Convert to Word"):
            try:
                if PdfReader is None:
                    st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                    st.stop()
                with st.spinner("Converting..."):
                    reader = PdfReader(io.BytesIO(f.read()))
                    doc = Document()
                    for p in reader.pages:
                        txt = p.extract_text() or ""
                        doc.add_paragraph(txt)
                    out = io.BytesIO(); doc.save(out); out.seek(0)
                st.download_button("📥 Download .docx", out.getvalue(), file_name="converted.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            except Exception:
                st.error(traceback.format_exc())

def render_pdf_to_excel():
    st.markdown("---")
    st.markdown("### 📊 Konversi PDF ke Excel (Text per Halaman)")
    f = st.file_uploader("Upload PDF", type="pdf")
    if f and st.button("Convert to Excel (text)"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Converting..."):
                reader = PdfReader(io.BytesIO(f.read()))
                rows = []
                for i, p in enumerate(reader.pages):
                    rows.append({"page": i+1, "text": p.extract_text() or ""})
                df = pd.DataFrame(rows)
                excel_bytes = df_to_excel_bytes(df)
                st.download_button("📥 Download Excel", excel_bytes, file_name="pdf_text.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        except Exception:
            st.error(traceback.format_exc())

def render_encrypt_pdf():
    st.markdown("---")
    st.markdown("### 🔒 Kunci (Encrypt) PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    pw = st.text_input("Password", type="password")
    if f and pw and st.button("Encrypt"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Mengunci PDF..."):
                reader = PdfReader(io.BytesIO(f.read()))
                writer = PdfWriter()
                for p in reader.pages:
                    writer.add_page(p)
                try_encrypt(writer, pw)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download encrypted.pdf", buf.getvalue(), file_name="encrypted.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_decrypt_pdf():
    st.markdown("---")
    st.markdown("### 🔓 Buka Kunci (Decrypt) PDF")
    f = st.file_uploader("Upload encrypted PDF", type="pdf")
    pw = st.text_input("Password for decryption", type="password")
    if f and pw and st.button("Decrypt"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Membuka PDF..."):
                reader = PdfReader(io.BytesIO(f.read()))
                if getattr(reader, "is_encrypted", False):
                    reader.decrypt(pw)
                writer = PdfWriter()
                for p in reader.pages:
                    writer.add_page(p)
                buf = io.BytesIO(); writer.write(buf); buf.seek(0)
            st.download_button("📥 Download decrypted.pdf", buf.getvalue(), file_name="decrypted.pdf", mime="application/pdf")
        except Exception:
            st.error(traceback.format_exc())

def render_batch_lock():
    st.markdown("---")
    st.markdown("### 🔐 Batch Lock PDF Berdasarkan Daftar Excel")
    excel_file = st.file_uploader("Upload Excel (filename,password) or CSV", type=["xlsx","csv"])
    pdfs = st.file_uploader("Upload PDFs (multiple)", type="pdf", accept_multiple_files=True)
    if excel_file and pdfs and st.button("Batch Lock"):
        try:
            if PdfReader is None:
                st.error("PyPDF2 tidak terinstall (pip install PyPDF2)")
                st.stop()
            with st.spinner("Batch locking PDFs..."):
                if excel_file.name.lower().endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(excel_file.read()))
                else:
                    df = pd.read_excel(io.BytesIO(excel_file.read()))
                pdf_map = {p.name: p.read() for p in pdfs}
                out_map = {}
                not_found = []
                total = len(df)
                prog = st.progress(0)
                for idx, (_, row) in enumerate(df.iterrows()):
                    cols = [c.lower() for c in df.columns]
                    try:
                        target_col = next((c for c in df.columns if c.lower() in ('filename', 'nama_file')), None)
                        pwd_col = next((c for c in df.columns if c.lower() in ('password', 'kata_sandi')), None)

                        if target_col and pwd_col:
                            target = str(row[target_col]).strip()
                            pwd = str(row[pwd_col]).strip()
                        else:
                            target = str(row.iloc[0]).strip()
                            pwd = str(row.iloc[1]).strip()
                    except Exception:
                        target = None; pwd = None
                    
                    if target and pwd:
                        matches = [k for k in pdf_map.keys() if k == target]
                        if matches:
                            key = matches[0]
                            reader = PdfReader(io.BytesIO(pdf_map[key]))
                            writer = PdfWriter()
                            for p in reader.pages:
                                writer.add_page(p)
                            try_encrypt(writer, pwd)
                            b = io.BytesIO(); writer.write(b); out_map[f"locked_{key}"] = b.getvalue()
                        else:
                            not_found.append(target)
                    prog.progress(int((idx+1)/total*100))
                if out_map:
                    st.download_button("📥 Download locked_pdfs.zip", make_zip_from_map(out_map), file_name="locked_pdfs.zip", mime="application/zip")
                if not_found:
                    st.warning(f"{len(not_found)} files not found sample: {not_found[:10]}")
        except Exception:
            st.error(traceback.format_exc())

def render_preview_pdf():
    st.markdown("---")
    st.markdown("### 👁️ Preview PDF")
    f = st.file_uploader("Upload PDF", type="pdf")
    mode = st.radio("Preview mode", ["First page (fast)", "All pages (slow)"])
    if f and st.button("Show Preview"):
        try:
            if PdfReader is None and not PDF2IMAGE_AVAILABLE:
                st.error("PyPDF2/pdf2image tidak terinstall.")
                st.stop()
            with st.spinner("Preparing preview..."):
                pdf_bytes = f.read()
                if PDF2IMAGE_AVAILABLE:
                    if mode.startswith("First"):
                        if convert_from_bytes is not None:
                            imgs = convert_from_bytes(pdf_bytes, first_page=1, last_page=1)
                        else:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(pdf_bytes)
                            tmp_path = tmp.name
                            imgs = convert_from_path(tmp_path, first_page=1, last_page=1)
                            try:
                                os.unlink(tmp_path)
                            except Exception:
                                pass
                        buf = io.BytesIO(); imgs[0].save(buf, format="PNG"); buf.seek(0)
                        st.image(buf.getvalue(), caption="Page 1")
                    else:
                        if convert_from_bytes is not None:
                            images = convert_from_bytes(pdf_bytes)
                        else:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                                tmp.write(pdf_bytes)
                            tmp_path = tmp.name
                            images = convert_from_path(tmp_path)
                            try:
                                os.unlink(tmp_path)
                            except:
                                pass
                        for i, img in enumerate(images):
                            buf = io.BytesIO(); img.save(buf, format="PNG"); st.image(buf.getvalue(), caption=f"Page {i+1}")
                else:
                    reader = PdfReader(io.BytesIO(pdf_bytes))
                    if mode.startswith("First"):
                        text = reader.pages[0].extract_text() or "Teks tidak dapat diekstrak (Mungkin gambar)"
                        st.text_area("Preview Page 1 (Text only)", text, height=300)
                    else:
                        all_text = [p.extract_text() or "Teks tidak dapat diekstrak" for p in reader.pages]
                        st.text_area("Preview All Pages (Text only)", "\n\n--- Halaman Berikutnya ---\n\n".join(all_text), height=500)
        except Exception as e:
            st.error(f"Gagal menampilkan preview. Pastikan Poppler terinstall untuk konversi ke gambar. Error: {e}")
            traceback.print_exc()

# -------------- File Tools --------------
def render_file_tools():
    add_back_to_dashboard_button() 
    st.subheader("📁 File Tools")

    file_tool = st.selectbox("Pilih Fitur File", [
        "📦 Zip / Unzip File", 
        "🔄 Konversi Dasar (misal: TXT/CSV/JSON -> Excel)", 
        "🔍 Cek Keberadaan Library"
    ])

    if file_tool == "📦 Zip / Unzip File":
        render_zip_unzip()
    elif file_tool == "🔄 Konversi Dasar (misal: TXT/CSV/JSON -> Excel)":
        render_basic_conversion()
    elif file_tool == "🔍 Cek Keberadaan Library":
        display_library_status()

def render_zip_unzip():
    st.markdown("---")
    st.subheader("Kompres ke ZIP atau Ekstrak dari ZIP")
    mode = st.radio("Pilih Mode", ["Compress to ZIP", "Extract from ZIP"])

    if mode == "Compress to ZIP":
        files = st.file_uploader("Unggah File (Multiple)", accept_multiple_files=True)
        if files and st.button("Buat ZIP"):
            try:
                out_map = {f.name: f.read() for f in files}
                zipb = make_zip_from_map(out_map)
                st.download_button("📥 Unduh ZIP", zipb, file_name="compressed_files.zip", mime="application/zip")
                st.success("✅ Kompresi selesai.")
            except Exception as e:
                st.error(f"Gagal: {e}")

    elif mode == "Extract from ZIP":
        f = st.file_uploader("Unggah File ZIP", type=["zip"])
        if f and st.button("Ekstrak ke Folder/ZIP"):
            try:
                z = zipfile.ZipFile(io.BytesIO(f.read()))
                extracted_files = {}
                for name in z.namelist():
                    if not name.endswith('/'):
                        extracted_files[name] = z.read(name)
                
                if extracted_files:
                    st.download_button("📥 Unduh Hasil Ekstraksi (ZIP)", make_zip_from_map(extracted_files), file_name="extracted_content.zip", mime="application/zip")
                    st.info(f"✅ {len(extracted_files)} file berhasil diekstrak.")
                else:
                    st.warning("File ZIP kosong atau hanya berisi folder.")
            except Exception as e:
                st.error(f"Gagal ekstrak: {e}")

def render_basic_conversion():
    st.markdown("---")
    st.subheader("Konversi Data ke Excel")
    f = st.file_uploader("Unggah file (TXT, CSV, JSON)", type=["txt", "csv", "json"])
    if f:
        df = None
        try:
            if f.name.lower().endswith(".csv"):
                df = pd.read_csv(io.BytesIO(f.read()))
            elif f.name.lower().endswith(".json"):
                df = pd.read_json(io.BytesIO(f.read()))
            elif f.name.lower().endswith(".txt"):
                df = pd.read_csv(io.BytesIO(f.read())) 
            
            if df is not None:
                st.dataframe(df.head())
                if st.button("Konversi ke Excel"):
                    excel_bytes = df_to_excel_bytes(df)
                    st.download_button("📥 Unduh Excel", excel_bytes, file_name="converted_file.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    st.success("✅ Konversi berhasil.")
        except Exception as e:
            st.error(f"Gagal memproses file: {e}")

# -------------- MCU Tools --------------
def render_mcu_tools():
    add_back_to_dashboard_button() 
    st.subheader("🩺 MCU Tools (Analisis Data Kesehatan)")
    st.warning("Fitur ini membutuhkan template Excel/PDF khusus untuk analisis. Pastikan format input data Anda sesuai.")
    
    mcu_tool = st.selectbox("Pilih Fitur MCU", [
        "📊 Organise by Excel (Original Logic) - Fitur Baru", 
        "📈 Dashboard Analisis Data MCU (Excel) - Diperbarui",
        "🔄 Konversi Laporan MCU (PDF) ke Data", 
    ], index=1)
    
    if mcu_tool == "📊 Organise by Excel (Original Logic) - Fitur Baru":
        render_organise_by_excel()
    elif mcu_tool == "📈 Dashboard Analisis Data MCU (Excel) - Diperbarui":
        render_mcu_dashboard()
    elif mcu_tool == "🔄 Konversi Laporan MCU (PDF) ke Data":
        render_pdf_to_data()

def render_organise_by_excel():
    st.markdown("---")
    st.subheader("📊 Organise by Excel (Original Logic)")
    st.info("Fitur ini akan membuat struktur folder di dalam file ZIP berdasarkan data Excel dan nama file PDF yang diunggah.")
    
    excel_up = st.file_uploader("Upload Excel (No_MCU, Nama, Departemen, JABATAN) or (filename,target_folder)", type=["xlsx","csv"], key="mcu_organize_excel")
    pdfs = st.file_uploader("Upload PDF files (multiple)", type="pdf", accept_multiple_files=True, key="mcu_organize_pdf")
    
    if excel_up and pdfs and st.button("Process MCU"):
        try:
            with st.spinner("Memproses MCU..."):
                # Baca Excel/CSV
                if excel_up.name.lower().endswith(".csv"):
                    df = pd.read_csv(io.BytesIO(excel_up.read()))
                else:
                    df = pd.read_excel(io.BytesIO(excel_up.read()))
                    
                pdf_map = {p.name: p.read() for p in pdfs}
                out_map = {}
                not_found = []
                
                # Logika Organise by Excel
                if all(c in df.columns for c in ["No_MCU","Nama","Departemen","JABATAN"]):
                    st.info("Mode: Organisasi berdasarkan kolom **No_MCU, Departemen, JABATAN** (Struktur: Dept/Jabatan/File.pdf).")
                    total = len(df)
                    prog = st.progress(0)
                    for idx, r in df.iterrows():
                        no = str(r["No_MCU"]).strip()
                        dept = clean_folder_name(r["Departemen"]) if not pd.isna(r["Departemen"]) else "Unknown_Dept"
                        jab = clean_folder_name(r["JABATAN"]) if not pd.isna(r["JABATAN"]) else "Unknown_JABATAN"
                        
                        matches = [k for k in pdf_map.keys() if k.startswith(no)] 
                        
                        if matches:
                            out_map[f"{dept}/{jab}/{matches[0]}"] = pdf_map[matches[0]]
                        else:
                            not_found.append(no)
                        prog.progress(int((idx+1)/total*100))
                        
                elif "filename" in df.columns and "target_folder" in df.columns:
                    st.info("Mode: Organisasi berdasarkan kolom **filename** dan **target_folder** (Struktur: Folder/File.pdf).")
                    total = len(df)
                    prog = st.progress(0)
                    for idx, r in df.iterrows():
                        fn = str(r["filename"]).strip()
                        tgt = clean_folder_name(r["target_folder"])
                        
                        if fn in pdf_map:
                            out_map[f"{tgt}/{fn}"] = pdf_map[fn]
                        else:
                            not_found.append(fn)
                        prog.progress(int((idx+1)/total*100))
                        
                else:
                    st.error("Format Excel/CSV tidak valid. Diperlukan kolom: **No_MCU, Nama, Departemen, JABATAN** ATAU **filename, target_folder**.")
                    
            # Hasil Download
            if out_map:
                zipb = make_zip_from_map(out_map)
                st.download_button("📥 Download MCU zip", zipb, file_name="mcu_structured.zip", mime="application/zip")
                st.success(f"✅ {len(out_map)} file berhasil diproses dan diatur strukturnya.")
            else:
                st.warning("Tidak ada file yang berhasil diproses.")
                
            if not_found:
                st.warning(f"{len(not_found)} ID/File tidak ditemukan di file PDF yang diunggah. Contoh: {not_found[:10]}")
                
        except Exception:
            st.error(f"Terjadi kesalahan saat memproses data. Cek format Excel Anda: {traceback.format_exc()}")

def render_mcu_dashboard():
    st.markdown("---")
    st.subheader("📈 Dashboard Analisis Hasil MCU Massal (Diperbarui)")
    st.markdown("Unggah data hasil MCU (Excel/CSV) untuk analisis cepat, visualisasi, dan filter data.")
    
    uploaded_file = st.file_uploader(
        "Unggah file Data MCU (Excel/CSV):",
        type=["xlsx", "csv"],
        key="mcu_data_uploader_new"
    )

    if uploaded_file:
        try:
            # 1. Baca File
            with st.spinner("Membaca data dan normalisasi kolom..."):
                if uploaded_file.name.lower().endswith('.csv'):
                    df = pd.read_csv(io.BytesIO(uploaded_file.read()))
                else:
                    df = pd.read_excel(io.BytesIO(uploaded_file.read()))

                st.success(f"✅ Data berhasil dimuat. Total Baris: {len(df)}")
                
                # Normalisasi kolom
                df = clean_column_names(df)
            
            st.markdown("#### Preview Data (5 Baris Teratas)")
            st.dataframe(df.head(), use_container_width=True)

            st.markdown("---")
            st.markdown("### 📊 Visualisasi & Analisis Cepat Status")
            
            # 2. Analisis Status Kesehatan
            status_cols = [col for col in df.columns if 'status' in col or 'fit' in col or 'hasil' in col]
            
            if status_cols:
                col1, col2 = st.columns([2, 1])
                with col1:
                    status_col = st.selectbox(
                        "Pilih Kolom Utama Status/Hasil:", 
                        status_cols, 
                        index=0, 
                        key="select_status_col",
                        help="Pilih kolom yang berisi status akhir MCU (misal: fit, unfit)."
                    )
                
                st.markdown(f"##### 1. Distribusi Status Kesehatan (`{status_col}`)")
                
                # Data Cleaning & Aggregation
                df[status_col] = df[status_col].astype(str).str.strip().str.upper().fillna("TIDAK DIKETAHUI")
                status_counts = df[status_col].value_counts().reset_index()
                status_counts.columns = [status_col, 'Jumlah']
                status_counts = status_counts.sort_values(by='Jumlah', ascending=False)
                
                if len(status_counts) > 0:
                    st.dataframe(status_counts, use_container_width=True)
                    st.bar_chart(status_counts.set_index(status_col))
                    
                    # Download Data Agregat
                    excel_bytes = df_to_excel_bytes(status_counts)
                    st.download_button(
                        "📥 Unduh Data Agregasi Status (Excel)", 
                        data=excel_bytes, 
                        file_name="status_agregat.xlsx", 
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                else:
                    st.info("Kolom status/hasil tidak memiliki data unik yang valid.")
            else:
                st.warning("Kolom yang mengandung kata 'status', 'fit', atau 'hasil' tidak ditemukan untuk Analisis Cepat Status. Cek penamaan kolom Anda.")

            st.markdown("---")
            
            # 3. Analisis Categorical/Filter
            st.markdown("### 📊 Filter dan Analisis Data Kategorikal")
            
            filter_cols = [
                col for col in df.columns 
                if df[col].dtype == 'object' and df[col].nunique() > 1 and df[col].nunique() <= 50
            ]
            
            if filter_cols:
                col_to_analyze = st.selectbox(
                    "Pilih Kolom Kategorikal (misal: Departemen, Jabatan, Gender):", 
                    filter_cols,
                    key="select_filter_col"
                )
                
                st.write(f"##### Distribusi Nilai dalam Kolom `{col_to_analyze}`")
                
                # Data Cleaning & Aggregation
                df[col_to_analyze] = df[col_to_analyze].astype(str).str.strip().str.upper().fillna("TIDAK DIKETAHUI")
                cat_counts = df[col_to_analyze].value_counts().reset_index()
                cat_counts.columns = [col_to_analyze, 'Jumlah']
                
                st.dataframe(cat_counts, use_container_width=True)
                st.bar_chart(cat_counts.set_index(col_to_analyze))
                
                # Download Agregat Kategorikal
                excel_bytes_cat = df_to_excel_bytes(cat_counts)
                st.download_button(
                    f"📥 Unduh Data Agregasi {col_to_analyze} (Excel)", 
                    data=excel_bytes_cat, 
                    file_name=f"{col_to_analyze.lower()}_agregat.xlsx", 
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
                
                # Filter Data Mentah
                st.markdown("---")
                st.markdown("##### Tampilkan Data Mentah Terfilter")
                
                filter_values = ["-- Pilih Nilai untuk Filter Data --"] + list(df[col_to_analyze].unique())
                selected_value = st.selectbox(
                    f"Pilih Nilai `{col_to_analyze}` untuk Menampilkan Data:", 
                    filter_values,
                    key="select_cat_value"
                )
                
                if selected_value != "-- Pilih Nilai untuk Filter Data --":
                    df_filtered = df[df[col_to_analyze] == selected_value]
                    st.info(f"Menampilkan **{len(df_filtered)}** baris data untuk `{selected_value}`.")
                    st.dataframe(df_filtered, use_container_width=True)
                    
                    # Download Filtered Data Mentah
                    excel_bytes_filtered = df_to_excel_bytes(df_filtered)
                    st.download_button(
                        f"📥 Unduh Data Filtered ({selected_value}) (Excel)", 
                        data=excel_bytes_filtered, 
                        file_name=f"data_filtered_{selected_value.lower().replace(' ', '_')}.xlsx", 
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                    
                else:
                    st.info("Pilih nilai di atas untuk menampilkan dan mengunduh data mentah yang terfilter.")
            else:
                st.info("Tidak ada kolom kategorikal yang cocok (object type dengan 2-50 nilai unik) untuk analisis mendalam. Pastikan kolom seperti 'Departemen' atau 'Gender' bertipe teks.")

        except Exception as e:
            st.error(f"Gagal memuat atau memproses file. Pastikan file Excel/CSV Anda valid: {e}")
            traceback.print_exc()

def render_pdf_to_data():
    st.markdown("---")
    st.subheader("🔄 Ekstraksi Data dari Laporan MCU PDF")
    st.warning("Fitur ini sangat bergantung pada struktur dan format PDF. Mungkin memerlukan konfigurasi kustom.")
    
    pdf_up = st.file_uploader("Unggah Laporan MCU PDF:", type=["pdf"], key="mcu_pdf_up")
    
    if pdf_up and st.button("Ekstrak Data"):
        st.error("Fitur ini adalah placeholder dan memerlukan logic ekstraksi PDF yang kompleks untuk diimplementasikan.")

# -------------- Tentang --------------
def render_about():
    add_back_to_dashboard_button() 
    st.subheader("💡 Tentang Master App – Tools MCU")
    st.markdown("""
    **KAY App** adalah aplikasi serbaguna berbasis Streamlit untuk membantu:
    -  **🖼️ Kompres Foto & Gambar**
    -  **📄 Pengelolaan Dokumen PDF** (gabung, pisah, proteksi, ekstraksi, Reorder/Hapus Halaman, Batch Rename, **Terjemahan**)
    -  **🩺 Analisis & Pengolahan Hasil MCU** (Dashboard Analisis Data, **Organise by Excel**)
    -  **📁 Manajemen File & Konversi Dasar** (Batch Rename/Format Gambar, Batch Rename PDF)

    <br>
    
    ### Kebutuhan Library Tambahan
    Beberapa fitur memerlukan library tambahan (instal di environment Anda):
    - `PyPDF2` (Dasar PDF): `pip install PyPDF2`
    - `pdfplumber` untuk ekstraksi tabel teks: `pip install pdfplumber`
    - `python-docx` untuk menghasilkan .docx: `pip install python-docx`
    - `deep-translator` untuk fitur terjemahan PDF: `pip install deep-translator`
    - `pdf2image` + poppler untuk konversi PDF->Gambar / Preview gambar: `pip install pdf2image`
    - `pandas` & `openpyxl` untuk Analisis MCU dan Batch Rename by Excel: `pip install pandas openpyxl`
    """)
    st.info("Data diproses di server tempat Streamlit dijalankan. Untuk mengaktifkan semua fitur, pasang dependensi yang diperlukan.")

# ----------------- Footer -----------------
def render_footer():
    st.markdown("""
    <hr style="border: none; border-top: 1px solid #cfe2ff; margin-top: 1.5rem;">
    <p style="text-align: right; color: #a0a0a0; font-size: 0.8rem;">
    Developed by AR - 2025
    </p>
    """, unsafe_allow_html=True)

# ----------------- Main Application Logic -----------------
def main():
    # Render the appropriate page based on menu selection
    if menu == "Dashboard":
        render_dashboard()
    elif menu == "Kompres Foto":
        render_compress_foto()
    elif menu == "PDF Tools":
        render_pdf_tools()
    elif menu == "MCU Tools":
        render_mcu_tools()
    elif menu == "File Tools":
        render_file_tools()
    elif menu == "Tentang":
        render_about()
    
    # Always render the footer
    render_footer()

if __name__ == "__main__":
    main()
